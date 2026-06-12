"""Load knowledge documents from disk."""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def load_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    raise ValueError(f"Unsupported document type: {path}")


def load_knowledge_base(directory: Path) -> list[Document]:
    docs: list[Document] = []
    for path in sorted(directory.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            content = load_document(path)
            if content.strip():
                docs.append(
                    Document(
                        page_content=content,
                        metadata={
                            "source": path.name,
                            "path": str(path),
                            "type": path.suffix.lower(),
                        },
                    )
                )
    return docs
